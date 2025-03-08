"""
utils/create_reference_template.py - Creates Word document template for pandoc

This module creates a reference Word document defining styles and formatting
for pandoc to use when converting Markdown documents to Word format.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_reference_template(output_path: str = None) -> str:
    """
    Create a reference Word document that defines styles for pandoc conversion.
    
    Args:
        output_path: Optional custom path for the reference document.
                    If not provided, uses templates/reference.docx
                    
    Returns:
        Path to the created reference document
        
    Raises:
        IOError: If unable to create the document or directory
    """
    # Use default path if none provided
    if output_path is None:
        # Ensure templates directory exists
        templates_dir = "templates"
        if not os.path.exists(templates_dir):
            try:
                os.makedirs(templates_dir)
                print(f"Created templates directory at {templates_dir}")
            except OSError as e:
                raise IOError(f"Failed to create templates directory: {str(e)}")
                
        output_path = os.path.join(templates_dir, "reference.docx")
    
    # Create a new Word document
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Define heading styles
    for level in range(1, 4):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            heading_style = doc.styles[style_name]
            font = heading_style.font
            font.name = 'Arial'
            font.size = Pt(16 - (level * 2))  # H1=16pt, H2=14pt, H3=12pt
            font.bold = True
            if level == 1:
                font.color.rgb = RGBColor(0, 0, 0)  # Black for H1
            
            # Set paragraph formatting
            paragraph_format = heading_style.paragraph_format
            paragraph_format.space_before = Pt(12)
            paragraph_format.space_after = Pt(6)
            paragraph_format.keep_with_next = True
    
    # Normal text style
    if 'Normal' in doc.styles:
        normal_style = doc.styles['Normal']
        font = normal_style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Set paragraph formatting
        paragraph_format = normal_style.paragraph_format
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.15
    
    # Create a custom style for the right-aligned header
    try:
        right_align_style = doc.styles.add_style('RightAligned', WD_STYLE_TYPE.PARAGRAPH)
        right_align_style.base_style = doc.styles['Normal']
        right_align_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        font = right_align_style.font
        font.bold = True
    except:
        # Style might already exist, not a critical error
        pass
    
    # Create example content to demonstrate styles
    # Header
    doc.add_paragraph("Example Header", style='RightAligned')
    
    # Headings with content
    doc.add_heading('Heading 1 Example', level=1)
    doc.add_paragraph('This is normal text under heading 1.')
    
    doc.add_heading('Heading 2 Example', level=2)
    doc.add_paragraph('This is normal text under heading 2.')
    
    doc.add_heading('Heading 3 Example', level=3)
    p = doc.add_paragraph('This is normal text with a ')
    p.add_run('bold').bold = True
    p.add_run(' and an ')
    p.add_run('italic').italic = True
    p.add_run(' word.')
    
    # Save the document
    try:
        doc.save(output_path)
        print(f"Created reference document at {output_path}")
    except Exception as e:
        raise IOError(f"Failed to save reference document: {str(e)}")
    
    return output_path

if __name__ == "__main__":
    create_reference_template()